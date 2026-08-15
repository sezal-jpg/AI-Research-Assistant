from pathlib import Path
from docx import Document as DocxDocument
from langchain_core.documents import Document 
from app.core.logger import logger

class DocxLoader:
    
    def load(self,file_path:Path):
        
        logger.info(f'Loading DOCX: {file_path.name}')
        document=DocxDocument(file_path)
        docs=[]
        
        for paragraph in document.paragraphs:
            text=paragraph.text.strip()
            
            if text:
                 docs.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source_file": file_path.name,
                            "source_type": "docx",
                        }
                    )
                )
        for table_index,table in enumerate(document.tables):
            rows=[]
            
            for row in table.rows:
                cells=[cell.text.strip() for cell in row.cells]   
                rows.append("|".join(cells))
            table_text="\n".join(rows)
            
            if table_text.strip():
                docs.append(Document(page_content=table_text,metadata={
                    'source_file':file_path.name,
                    'source_type':'docx',
                    'table':table_index+1,
                }))    
        logger.info(f'Loaded {len(docs)} DOCX documents')    
        
        return docs
    
docx_loader=DocxLoader()              